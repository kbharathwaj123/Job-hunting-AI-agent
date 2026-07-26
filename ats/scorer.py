"""
Lightweight ATS simulator.

Real ATS systems (Workday, Taleo, Greenhouse's own parser, etc.) mostly do:
  1. Parse resume text (fails badly on tables/graphics/columns/text-boxes —
     hence: keep resume as clean single-column text).
  2. Keyword/skill match against the job description.
  3. Sometimes basic section detection (Experience, Education, Skills).

This module approximates step 2 with keyword overlap + fuzzy matching,
enough to (a) score/rank jobs and (b) tell the LLM which keywords are
missing so it can honestly work them into your resume where you actually
have that experience.
"""

import re
from rapidfuzz import fuzz

STOPWORDS = {
    "and", "the", "a", "to", "of", "in", "for", "with", "on", "is", "are",
    "as", "an", "or", "be", "will", "you", "your", "we", "our", "at", "by",
}


def extract_keywords(text: str) -> set:
    words = re.findall(r"[A-Za-z][A-Za-z0-9+#.]{1,}", text.lower())
    return {w for w in words if w not in STOPWORDS and len(w) > 2}


def score_resume_against_job(resume_text: str, job_description: str):
    """
    Returns (score 0-100, matched_keywords, missing_keywords)
    """
    job_kw = extract_keywords(job_description)
    resume_kw = extract_keywords(resume_text)

    matched = set()
    missing = set()

    for kw in job_kw:
        # fuzzy match to catch e.g. "Selenium" vs "Selenium WebDriver"
        best = max((fuzz.partial_ratio(kw, rkw) for rkw in resume_kw), default=0)
        if best >= 85:
            matched.add(kw)
        else:
            missing.add(kw)

    score = round(100 * len(matched) / max(len(job_kw), 1), 1)
    return score, matched, missing


def format_check(docx_path: str) -> list:
    """
    Basic ATS-safety check on the resume file itself.
    Flags things that commonly break real ATS parsers.
    """
    from docx import Document
    issues = []
    doc = Document(docx_path)

    if len(doc.tables) > 0:
        issues.append("Resume contains tables — many ATS parsers scramble table content. Prefer plain text sections.")

    for section in doc.sections:
        pass  # placeholder for header/footer text-box checks if needed

    if not doc.paragraphs or len(doc.paragraphs) < 5:
        issues.append("Very little parsable text detected — check for text stored inside images/text boxes.")

    return issues
